import os
import re
import logging
import shutil
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from urllib.parse import urlparse
import validators
import torch
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Depends, Security, Request, Response
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from supabase import create_client, Client
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS as LC_FAISS
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig
import fitz
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import requests
from tqdm import tqdm
import uvicorn
from dotenv import load_dotenv

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# FastAPI setup
app = FastAPI(title="RAG Backend API")

# CORS configuration for frontend compatibility
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "https://*.ngrok.io",
        "https://a27adc90-f8a2-4041-bb28-3bfff1842b21.lovableproject.com"
    ],
    allow_credentials=True,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# Custom middleware to handle OPTIONS requests
@app.middleware("http")
async def handle_options_requests(request: Request, call_next):
    if request.method == "OPTIONS":
        logger.info(f"Handling OPTIONS request for {request.url.path}")
        response = Response(status_code=200)
        response.headers["Access-Control-Allow-Origin"] = request.headers.get("origin", "*")
        response.headers["Access-Control-Allow-Methods"] = "GET,POST,DELETE,OPTIONS"
        response.headers["Access-Control-Allow-Headers"] = "*"
        response.headers["Access-Control-Allow-Credentials"] = "true"
        return response
    return await call_next(request)

# Load environment variables
load_dotenv()

# Supabase configuration
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
if not SUPABASE_URL or not SUPABASE_KEY:
    logger.error("SUPABASE_URL and SUPABASE_KEY must be set in environment variables")
    raise ValueError("SUPABASE_URL and SUPABASE_KEY must be set in environment variables")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Authentication setup
security = HTTPBearer()

async def get_current_user(credentials: HTTPAuthorizationCredentials = Security(security)):
    try:
        token = credentials.credentials
        user = supabase.auth.get_user(token)
        if not user or not user.user:
            raise HTTPException(status_code=401, detail="Invalid or expired token")
        supabase.postgrest.auth(token)
        if not user.user.id:
            raise HTTPException(status_code=401, detail="User ID is missing in authentication response")
        return user.user
    except Exception as e:
        logger.error(f"Authentication failed: {str(e)}")
        raise HTTPException(status_code=401, detail=f"Authentication failed: {str(e)}")

# Document loading functions
def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    return text.strip()

def load_pdf(file_path: str) -> List[dict]:
    try:
        doc = fitz.open(file_path)
        texts = []
        for i, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                texts.append({
                    "text": clean_text(page_text),
                    "metadata": {
                        "source": os.path.basename(file_path),
                        "type": "pdf",
                        "page": i + 1
                    }
                })
        doc.close()
        return texts
    except Exception as e:
        logger.error(f"Failed to process PDF {file_path}: {str(e)}")
        raise ValueError(f"Failed to process PDF {file_path}: {str(e)}")

def load_text_file(file_path: str) -> List[dict]:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            raw_text = f.read()
        return [{
            "text": clean_text(raw_text),
            "metadata": {
                "source": os.path.basename(file_path),
                "type": "txt"
            }
        }]
    except Exception as e:
        logger.error(f"Failed to process text file {file_path}: {str(e)}")
        raise ValueError(f"Failed to process text file {file_path}: {str(e)}")

def load_docx_file(file_path: str) -> List[dict]:
    try:
        doc = DocxDocument(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [{
            "text": clean_text(full_text),
            "metadata": {
                "source": os.path.basename(file_path),
                "type": "docx"
            }
        }]
    except Exception as e:
        logger.error(f"Failed to process DOCX file {file_path}: {str(e)}")
        raise ValueError(f"Failed to process DOCX file {file_path}: {str(e)}")

def load_webpage(url: str) -> List[dict]:
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        raw_text = soup.get_text(separator="\n")
        return [{
            "text": clean_text(raw_text),
            "metadata": {
                "source": url,
                "type": "web"
            }
        }]
    except Exception as e:
        logger.error(f"Failed to process webpage {url}: {str(e)}")
        raise ValueError(f"Failed to process webpage {url}: {str(e)}")

# Chunking and embedding
def process_documents(user_id: str, web_urls: List[str] = []) -> LC_FAISS:
    user_dir = Path("input_docs") / user_id
    user_dir.mkdir(parents=True, exist_ok=True)
    docs = []
    for filename in os.listdir(user_dir):
        filepath = os.path.join(user_dir, filename)
        if filename.lower().endswith(".pdf"):
            docs.extend(load_pdf(filepath))
        elif filename.lower().endswith(".txt"):
            docs.extend(load_text_file(filepath))
        elif filename.lower().endswith(".docx"):
            docs.extend(load_docx_file(filepath))
    for url in web_urls:
        if validators.url(url):
            docs.extend(load_webpage(url))
        else:
            logger.warning(f"Invalid URL skipped: {url}")

    if not docs:
        raise ValueError("No valid documents or URLs found to process")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    chunked_docs = []
    for doc in tqdm(docs, desc="Chunking documents"):
        if not doc["text"].strip():
            continue
        splits = text_splitter.split_text(doc["text"])
        for i, chunk in enumerate(splits):
            chunked_docs.append(Document(
                page_content=chunk,
                metadata={**doc["metadata"], "chunk": i + 1, "user_id": user_id}
            ))

    embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vector_store_path = Path("vector_store") / user_id
    vector_store_path.mkdir(parents=True, exist_ok=True)

    try:
        vector_store = LC_FAISS.from_documents(chunked_docs, embedding_model)
        vector_store.save_local(vector_store_path)
        logger.info(f"Vector store saved for user {user_id} at {vector_store_path}")
        return vector_store
    except Exception as e:
        logger.error(f"Failed to create vector store for user {user_id}: {str(e)}")
        raise ValueError(f"Failed to create vector store: {str(e)}")

# RAG setup with Zephyr-7B
def setup_rag(user_id: str):
    try:
        embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vector_store_path = Path("vector_store") / user_id
        if not vector_store_path.exists():
            raise ValueError(f"No vector store found for user {user_id}")
        retriever = LC_FAISS.load_local(
            vector_store_path,
            embedding_model,
            allow_dangerous_deserialization=True
        ).as_retriever(search_kwargs={"k": 5})

        logger.info("Using device: cpu for model loading")
        bnb_config = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_use_double_quant=True,
            bnb_4bit_quant_type="nf4",
            llm_int8_skip_modules=None
        )
        tokenizer = AutoTokenizer.from_pretrained("HuggingFaceH4/zephyr-7b-beta")
        model = AutoModelForCausalLM.from_pretrained(
            "HuggingFaceH4/zephyr-7b-beta",
            device_map=None,  # Explicitly use CPU
            quantization_config=bnb_config,
            torch_dtype=torch.float16,
            trust_remote_code=True
        )
        logger.info("Loaded Zephyr-7B with 4-bit quantization on CPU")
        logger.info(f"Model device: {next(model.parameters()).device}")

        generator = pipeline(
            "text-generation",
            model=model,
            tokenizer=tokenizer,
            max_new_tokens=512,
            do_sample=False
        )
        return retriever, generator
    except Exception as e:
        logger.error(f"Failed to initialize RAG components for user {user_id}: {str(e)}")
        raise ValueError(f"Failed to initialize RAG components: {str(e)}")

# Question answering
def ask_question(query: str, retriever, generator) -> tuple:
    if not query.strip():
        raise ValueError("Query cannot be empty")
    retrieved_docs = retriever.invoke(query)
    if not retrieved_docs:
        raise ValueError("No relevant documents found for the query")
    
    context = "\n\n".join([doc.page_content for doc in retrieved_docs])
    system_prompt = (
        "You are a knowledgeable and trustworthy assistant who only uses the given context to answer."
        " If the answer is not in the context, say 'I don't know'. Avoid hallucinating facts."
        " Be concise, clear, and include relevant details if needed."
    )
    prompt = (
        f"<|system|>\n{system_prompt}</s>\n"
        f"<|user|>\nContext:\n{context}\n\nQuestion: {query}</s>\n"
        f"<|assistant|>"
    )
    if len(prompt) > 4000:
        prompt = prompt[-4000:]

    try:
        raw_output = generator(prompt, max_new_tokens=512, do_sample=False)[0]["generated_text"]
        answer = raw_output.split("<|assistant|>")[-1].strip()
        sources = list({doc.metadata.get("source", "Unknown") for doc in retrieved_docs})
        return answer, sources
    except Exception as e:
        logger.error(f"Error generating answer for query '{query}': {str(e)}")
        raise ValueError(f"Failed to generate answer: {str(e)}")

# Pydantic models for request/response
class QueryRequest(BaseModel):
    question: str

class QueryResponse(BaseModel):
    answer: str
    sources: List[str]

# User-specific RAG components
user_rag_components = {}

# Cleanup utility
def cleanup_user_data(user_id: str):
    try:
        user_dir = Path("input_docs") / user_id
        vector_store_path = Path("vector_store") / user_id
        if user_dir.exists():
            shutil.rmtree(user_dir)
            logger.info(f"Cleaned up user directory: {user_dir}")
        if vector_store_path.exists():
            shutil.rmtree(vector_store_path)
            logger.info(f"Cleaned up vector store: {vector_store_path}")
        if user_id in user_rag_components:
            del user_rag_components[user_id]
            logger.info(f"Removed RAG components for user {user_id}")
    except Exception as e:
        logger.error(f"Error cleaning up user data for {user_id}: {str(e)}")

@app.post("/upload", response_model=dict)
async def upload_file(file: UploadFile = File(...), current_user=Depends(get_current_user)):
    try:
        if not file.filename.lower().endswith((".pdf", ".txt", ".docx")):
            raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, TXT, or DOCX allowed")
        
        user_dir = Path("input_docs") / current_user.id
        user_dir.mkdir(parents=True, exist_ok=True)
        file_path = user_dir / file.filename
        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
        
        file_type = file.filename.split('.')[-1].lower()
        try:
            supabase.table("documents").insert({
                "filename": file.filename,
                "file_type": file_type,
                "uploaded_at": datetime.utcnow().isoformat(),
                "user_id": str(current_user.id)
            }).execute()
        except Exception as e:
            logger.error(f"Supabase insert failed for user {current_user.id}: {str(e)}")
            if "violates row-level security policy" in str(e):
                raise HTTPException(status_code=403, detail="Failed to save document metadata: RLS policy violation. Ensure user_id is set and table permissions are correct.")
            raise HTTPException(status_code=403, detail=f"Failed to save document metadata due to permission error: {str(e)}")
        
        vector_store = process_documents(current_user.id)
        retriever, generator = setup_rag(current_user.id)
        user_rag_components[current_user.id] = {"vector_store": vector_store, "retriever": retriever, "generator": generator}
        logger.info(f"File {file.filename} uploaded and processed for user {current_user.id}")
        return {"message": f"File {file.filename} uploaded and processed successfully"}
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Upload failed for user {current_user.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/add_url", response_model=dict)
async def add_url(url: str = Form(...), current_user=Depends(get_current_user)):
    try:
        if not validators.url(url):
            raise HTTPException(status_code=400, detail="Invalid URL")
        
        url_name = urlparse(url).netloc or url
        try:
            supabase.table("documents").insert({
                "filename": url_name,
                "file_type": "web",
                "uploaded_at": datetime.utcnow().isoformat(),
                "user_id": str(current_user.id)
            }).execute()
        except Exception as e:
            logger.error(f"Supabase insert failed for user {current_user.id}: {str(e)}")
            if "violates row-level security policy" in str(e):
                raise HTTPException(status_code=403, detail="Failed to save URL metadata: RLS policy violation. Ensure user_id is set and table permissions are correct.")
            raise HTTPException(status_code=403, detail=f"Failed to save URL metadata due to permission error: {str(e)}")
        
        vector_store = process_documents(current_user.id, web_urls=[url])
        retriever, generator = setup_rag(current_user.id)
        user_rag_components[current_user.id] = {"vector_store": vector_store, "retriever": retriever, "generator": generator}
        logger.info(f"URL {url} processed for user {current_user.id}")
        return {"message": f"URL {url} processed successfully"}
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"URL processing failed for user {current_user.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"URL processing failed: {str(e)}")

@app.post("/query", response_model=QueryResponse)
async def query(request: QueryRequest, current_user=Depends(get_current_user)):
    try:
        if current_user.id not in user_rag_components:
            raise HTTPException(status_code=400, detail="RAG system not initialized. Upload documents or URLs first.")
        retriever = user_rag_components[current_user.id]["retriever"]
        generator = user_rag_components[current_user.id]["generator"]
        answer, sources = ask_question(request.question, retriever, generator)
        
        try:
            supabase.table("interactions").insert({
                "user_id": str(current_user.id),
                "query": request.question,
                "response": answer,
                "sources": sources,
                "timestamp": datetime.utcnow().isoformat()
            }).execute()
        except Exception as e:
            logger.error(f"Supabase insert failed for user {current_user.id}: {str(e)}")
            if "violates row-level security policy" in str(e):
                raise HTTPException(status_code=403, detail="Failed to save interaction: RLS policy violation. Ensure user_id is set and table permissions are correct.")
            raise HTTPException(status_code=403, detail=f"Failed to save interaction due to permission error: {str(e)}")
        
        logger.info(f"Query answered for user {current_user.id}: {request.question}")
        return QueryResponse(answer=answer, sources=sources)
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Query failed for user {current_user.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Query failed: {str(e)}")

@app.get("/documents", response_model=dict)
async def get_documents(current_user=Depends(get_current_user)):
    try:
        response = supabase.table("documents").select("*").eq("user_id", str(current_user.id)).execute()
        logger.info(f"Fetched documents for user {current_user.id}")
        return {"documents": response.data}
    except Exception as e:
        logger.error(f"Failed to fetch documents for user {current_user.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch documents: {str(e)}")

@app.get("/get_interactions", response_model=dict)
async def get_interactions(current_user=Depends(get_current_user), limit: int = 50, offset: int = 0):
    try:
        response = supabase.table("interactions").select("*").eq("user_id", str(current_user.id)).order("timestamp", desc=True).range(offset, offset + limit - 1).execute()
        logger.info(f"Fetched interactions for user {current_user.id}")
        return {"interactions": response.data}
    except Exception as e:
        logger.error(f"Failed to fetch interactions for user {current_user.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch interactions: {str(e)}")

@app.delete("/cleanup", response_model=dict)
async def cleanup(current_user=Depends(get_current_user)):
    try:
        cleanup_user_data(current_user.id)
        try:
            supabase.table("documents").delete().eq("user_id", str(current_user.id)).execute()
            supabase.table("interactions").delete().eq("user_id", str(current_user.id)).execute()
        except Exception as e:
            logger.error(f"Supabase delete failed for user {current_user.id}: {str(e)}")
            if "violates row-level security policy" in str(e):
                raise HTTPException(status_code=403, detail="Failed to delete data: RLS policy violation. Ensure DELETE policies are set for documents and interactions tables.")
            raise HTTPException(status_code=403, detail=f"Failed to delete data due to permission error: {str(e)}")
        logger.info(f"Cleaned up all data for user {current_user.id}")
        return {"message": "User data cleaned up successfully"}
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Cleanup failed for user {current_user.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Cleanup failed: {str(e)}")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)