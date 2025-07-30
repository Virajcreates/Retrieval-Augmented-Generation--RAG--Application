import os
import re
import numpy as np
import pandas as pd
from tqdm import tqdm
import nltk
nltk.download('punkt')
from docx import Document as DocxDocument
import fitz
from bs4 import BeautifulSoup
import requests
from sentence_transformers import SentenceTransformer
import faiss
from langchain_community.vectorstores import FAISS as LC_FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import torch
from transformers import pipeline
import warnings
warnings.filterwarnings("ignore")
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import shutil
from supabase import create_client, Client
from datetime import datetime
from typing import List
from urllib.parse import urlparse
import validators

# Supabase configuration
SUPABASE_URL = "https://oopywyrxbckcxkayodcq.supabase.co"  # Replace with your Supabase Project URL
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Im9vcHl3eXJ4YmNrY3hrYXlvZGNxIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NTI1NjI3OTQsImV4cCI6MjA2ODEzODc5NH0.OsCAspsrItC8VIdJuEJMZ3-XWeM8A_FFIErRVNDD7PQ"  # Replace with your Supabase Anon Key
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Document loading functions
def load_pdf(file_path: str) -> List[dict]:
    try:
        doc = fitz.open(file_path)
        texts = []
        for i, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                texts.append({
                    "text": clean_text(page_text),
                    "metadata": {
                        "source": os.path.basename(file_path),
                        "type": "pdf",
                        "page": i + 1
                    }
                })
        doc.close()
        return texts
    except Exception as e:
        raise ValueError(f"Failed to process PDF {file_path}: {str(e)}")

def load_text_file(file_path: str) -> List[dict]:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            raw_text = f.read()
        return [{
            "text": clean_text(raw_text),
            "metadata": {
                "source": os.path.basename(file_path),
                "type": "txt"
            }
        }]
    except Exception as e:
        raise ValueError(f"Failed to process text file {file_path}: {str(e)}")

def load_docx_file(file_path: str) -> List[dict]:
    try:
        doc = DocxDocument(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [{
            "text": clean_text(full_text),
            "metadata": {
                "source": os.path.basename(file_path),
                "type": "docx"
            }
        }]
    except Exception as e:
        raise ValueError(f"Failed to process DOCX file {file_path}: {str(e)}")

def load_webpage(url: str) -> List[dict]:
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        raw_text = soup.get_text(separator="\n")
        return [{
            "text": clean_text(raw_text),
            "metadata": {
                "source": url,
                "type": "web"
            }
        }]
    except Exception as e:
        raise ValueError(f"Failed to process webpage {url}: {str(e)}")

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    return text.strip()

# Chunking and embedding
def process_documents(input_dir: str, web_urls: List[str] = []) -> LC_FAISS:
    docs = []
    for filename in os.listdir(input_dir):
        filepath = os.path.join(input_dir, filename)
        if filename.lower().endswith(".pdf"):
            docs.extend(load_pdf(filepath))
        elif filename.lower().endswith(".txt"):
            docs.extend(load_text_file(filepath))
        elif filename.lower().endswith(".docx"):
            docs.extend(load_docx_file(filepath))
    for url in web_urls:
        docs.extend(load_webpage(url))
    
    if not docs:
        raise ValueError("No valid documents found to process")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    chunked_docs = []
    for doc in docs:
        splits = text_splitter.split_text(doc["text"])
        for i, chunk in enumerate(splits):
            chunked_docs.append(Document(
                page_content=chunk,
                metadata={**doc["metadata"], "chunk": i + 1}
            ))
    
    embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vector_store = LC_FAISS.from_documents(chunked_docs, embedding_model)
    vector_store.save_local("vector_store")
    return vector_store

# Question answering
def setup_rag():
    try:
        embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        retriever = LC_FAISS.load_local(
            "vector_store",
            embedding_model,
            allow_dangerous_deserialization=True
        ).as_retriever(search_kwargs={"k": 5})
        
        generator = pipeline(
            "text2text-generation",
            model="google/flan-t5-base",
            tokenizer="google/flan-t5-base",
            max_length=512,
            device=0 if torch.cuda.is_available() else -1
        )
        return retriever, generator
    except Exception as e:
        raise ValueError(f"Failed to initialize RAG components: {str(e)}")

def ask_question(query: str, retriever, generator) -> tuple:
    if not query.strip():
        raise ValueError("Query cannot be empty")
    retrieved_docs = retriever.get_relevant_documents(query)
    if not retrieved_docs:
        raise ValueError("No relevant documents found for the query")
    context = "\n\n".join([doc.page_content for doc in retrieved_docs])
    prompt = (
        f"Answer the following question using only the context below:\n\n"
        f"Context:\n{context}\n\n"
        f"Question: {query}\nAnswer:"
    )
    response = generator(prompt, do_sample=False)[0]['generated_text']
    answer = response.strip()
    sources = list({doc.metadata.get('source', 'Unknown') for doc in retrieved_docs})
    return answer, sources

# FastAPI setup
app = FastAPI(title="RAG Backend API")

# CORS configuration for Lovable frontend and local testing
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "https://*.ngrok.io", "https://a27adc90-f8a2-4041-bb28-3bfff1842b21.lovableproject.com"],  # Update with actual Lovable URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize RAG components
input_dir = "input_docs"
os.makedirs(input_dir, exist_ok=True)
os.makedirs("vector_store", exist_ok=True)
vector_store = None
retriever, generator = None, None

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Validate file type
        if not file.filename.lower().endswith((".pdf", ".txt", ".docx")):
            raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, TXT, or DOCX allowed")
        
        # Save file
        file_path = os.path.join(input_dir, file.filename)
        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
        
        # Store metadata in Supabase
        file_type = file.filename.split('.')[-1].lower()
        supabase.table("documents").insert({
            "filename": file.filename,
            "file_type": file_type,
            "uploaded_at": datetime.utcnow().isoformat()
        }).execute()
        
        # Process documents and update vector store
        global vector_store, retriever, generator
        vector_store = process_documents(input_dir)
        retriever, generator = setup_rag()
        return JSONResponse(content={"message": f"File {file.filename} uploaded and processed successfully"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/add_url")
async def add_url(url: str = Form(...)):
    try:
        # Validate URL
        if not validators.url(url):
            raise HTTPException(status_code=400, detail="Invalid URL")
        
        # Store URL metadata in Supabase
        url_name = urlparse(url).netloc or url
        supabase.table("documents").insert({
            "filename": url_name,
            "file_type": "web",
            "uploaded_at": datetime.utcnow().isoformat()
        }).execute()
        
        # Process URL and update vector store
        global vector_store, retriever, generator
        vector_store = process_documents(input_dir, web_urls=[url])
        retriever, generator = setup_rag()
        return JSONResponse(content={"message": f"URL {url} processed successfully"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"URL processing failed: {str(e)}")

@app.get("/query")
async def query(question: str):
    if not retriever or not generator:
        raise HTTPException(status_code=400, detail="RAG system not initialized. Upload documents or URLs first.")
    try:
        answer, sources = ask_question(question, retriever, generator)
        return JSONResponse(content={"answer": answer, "sources": sources})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query failed: {str(e)}")

@app.get("/documents")
async def get_documents():
    try:
        response = supabase.table("documents").select("*").execute()
        return JSONResponse(content={"documents": response.data})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch documents: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)